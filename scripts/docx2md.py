#!/usr/bin/env python3
"""
DOCX to Markdown converter for Sloth-BookSmith-Den.
Converts Word documents to booksmith-compatible Markdown with:
- Heading style mapping
- Page break preservation
- Image extraction
- Table conversion
- YAML frontmatter generation
"""
import os
import sys
import re
import argparse
from datetime import datetime
from docx import Document
from docx.oxml.ns import qn


# ─── Style Mapping ─────────────────────────────────────────────────────────
# Maps both style names and style IDs to Markdown heading levels
STYLE_MAP = {
    # Standard English names
    'Heading 1': 1, 'heading 1': 1,
    'Heading 2': 2, 'heading 2': 2,
    'Heading 3': 3, 'heading 3': 3,
    'Heading 4': 4, 'heading 4': 4,
    'Heading 5': 5, 'heading 5': 5,
    'Heading 6': 6, 'heading 6': 6,
    # Chinese Word names
    '标题 1': 1, '标题1': 1,
    '标题 2': 2, '标题2': 2,
    '标题 3': 3, '标题3': 3,
    '标题 4': 4, '标题4': 4,
    '标题 5': 5, '标题5': 5,
    '标题 6': 6, '标题6': 6,
    # Custom style IDs (from original document analysis)
    '21bc9c4b-6a32-43e5-beaa-fd2d792c5735': 1,
    '71e7dc79-1ff7-45e8-997d-0ebda3762b91': 2,
    'b63ee27f-4cf3-414c-9275-d88e3f90795e': 3,
}

SKIP_STYLES = {'Normal'}


def ensure_dir(path):
    os.makedirs(path, exist_ok=True)


def extract_images(doc, images_dir):
    """Extract all inline images from docx."""
    ensure_dir(images_dir)
    image_map = {}

    for i, shape in enumerate(doc.inline_shapes):
        try:
            blip = shape._inline.graphic.graphicData.pic.blipFill.blip
            rId = blip.embed
            image_part = doc.part.related_parts[rId]
            image_bytes = image_part.blob

            content_type = image_part.content_type
            if 'png' in content_type:
                ext = 'png'
            elif 'jpeg' in content_type or 'jpg' in content_type:
                ext = 'jpg'
            elif 'gif' in content_type:
                ext = 'gif'
            else:
                ext = 'png'

            filename = f'image_{i+1:03d}.{ext}'
            filepath = os.path.join(images_dir, filename)
            with open(filepath, 'wb') as f:
                f.write(image_bytes)
            image_map[i] = filename
            print(f'  Extracted: {filename} ({len(image_bytes)} bytes)')
        except Exception as e:
            print(f'  Failed to extract image {i}: {e}')

    return image_map


def has_page_break(para):
    """Check if paragraph contains a page break in any run."""
    for run in para.runs:
        xml = run._element.xml
        if 'w:type="page"' in xml or 'type="page"' in xml:
            return True
    return False


def process_run(run):
    """Convert a single run to markdown with bold/italic."""
    text = run.text or ''
    if not text:
        return ''

    bold = run.bold or (run.font.bold if run.font else False)
    italic = run.italic or (run.font.italic if run.font else False)

    # Escape markdown special chars
    text = text.replace('\\', '\\\\').replace('*', '\\*').replace('_', '\\_')

    if bold and italic:
        text = f'***{text}***'
    elif bold:
        text = f'**{text}**'
    elif italic:
        text = f'*{text}*'

    return text


def para_to_md(para, image_map, image_idx_ref):
    """Convert a paragraph to markdown."""
    style_name = para.style.name if para.style else 'None'
    style_id = para.style.style_id if para.style else ''
    text = para.text.strip()
    para_has_pb = has_page_break(para)

    # Skip empty Normal paragraphs (unless they have a page break)
    if style_name in SKIP_STYLES and not text and not para_has_pb:
        return ''

    # Heading — check both style name and style ID
    level = STYLE_MAP.get(style_name) or STYLE_MAP.get(style_id)
    if level:
        prefix = '#' * level
        return f"{prefix} {text}\n\n"

    # ds-markdown-paragraph - treat as blockquote
    if style_name == 'ds-markdown-paragraph':
        lines = para.text.split('\n')
        quoted = '\n'.join(f'> {line}' for line in lines if line.strip())
        return quoted + '\n\n'

    # Normal paragraphs (might have images or be empty)
    result = []
    for run in para.runs:
        xml = run._element.xml
        if 'w:type="page"' in xml or 'type="page"' in xml:
            accumulated = ''.join(result).strip()
            if accumulated:
                return accumulated + '\n\n<div class="page-break"></div>\n\n'
            else:
                return '<div class="page-break"></div>\n\n'
        result.append(process_run(run))

    combined = ''.join(result).strip()
    if combined:
        return combined + '\n\n'
    return ''


def table_to_md(table):
    """Convert a table to markdown."""
    rows = table.rows
    if not rows:
        return ''

    md_lines = []
    for i, row in enumerate(rows):
        cells = [cell.text.strip().replace('\n', ' ').replace('|', '\\|') for cell in row.cells]
        md_lines.append('| ' + ' | '.join(cells) + ' |')
        if i == 0:
            md_lines.append('|' + '|'.join(['---'] * len(cells)) + '|')

    return '\n'.join(md_lines) + '\n\n'


def build_frontmatter(doc, overrides=None):
    """Build YAML frontmatter from document properties."""
    cp = doc.core_properties
    overrides = overrides or {}

    title = overrides.get('title') or cp.title or ''
    if not title:
        # Fallback: use first Heading 1 paragraph as title
        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ''
            style_id = para.style.style_id if para.style else ''
            level = STYLE_MAP.get(style_name) or STYLE_MAP.get(style_id)
            if level == 1 and para.text.strip():
                title = para.text.strip()
                break

    author = overrides.get('author') or cp.author or ''
    subject = overrides.get('subtitle') or cp.subject or ''

    # Date
    date_val = overrides.get('date')
    if not date_val and cp.created:
        try:
            date_val = cp.created.strftime('%Y-%m-%d')
        except Exception:
            date_val = ''
    if not date_val:
        date_val = datetime.now().strftime('%Y-%m-%d')

    # Build YAML
    lines = ['---\n']
    if title:
        lines.append(f'title: "{title}"\n')
    if author:
        lines.append(f'author: "{author}"\n')
    if date_val:
        lines.append(f'date: "{date_val}"\n')
    if subject:
        lines.append(f'subtitle: "{subject}"\n')
    lines.append('---\n\n')
    return ''.join(lines)


def convert(input_docx, output_md=None, output_dir=None, images_dir=None,
            title=None, author=None, date=None, subtitle=None):
    """
    Convert a DOCX file to Markdown.

    Args:
        input_docx: Path to input .docx file
        output_md:  Path to output .md file (default: same dir/name as input)
        output_dir: Output directory (default: same as input)
        images_dir: Directory for extracted images (default: output_dir/images)
        title, author, date, subtitle: Overrides for YAML frontmatter

    Returns:
        Path to generated .md file
    """
    input_docx = os.path.abspath(input_docx)
    if not os.path.exists(input_docx):
        raise FileNotFoundError(f"Input file not found: {input_docx}")

    # Determine output paths
    if not output_dir:
        output_dir = os.path.dirname(input_docx) or os.getcwd()
    ensure_dir(output_dir)

    if not output_md:
        base = os.path.splitext(os.path.basename(input_docx))[0]
        output_md = os.path.join(output_dir, f"{base}.md")

    if not images_dir:
        images_dir = os.path.join(output_dir, 'images')
    ensure_dir(images_dir)

    print(f'Loading: {input_docx}')
    doc = Document(input_docx)

    # Extract images
    print('\nExtracting images...')
    image_map = extract_images(doc, images_dir)

    # Build frontmatter
    overrides = {}
    if title:
        overrides['title'] = title
    if author:
        overrides['author'] = author
    if date:
        overrides['date'] = date
    if subtitle:
        overrides['subtitle'] = subtitle

    md_parts = [build_frontmatter(doc, overrides)]

    # Process body
    print('\nConverting to Markdown...')
    image_idx_ref = [0]

    for element in doc.element.body:
        tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag

        if tag == 'p':
            for para in doc.paragraphs:
                if para._element is element:
                    md = para_to_md(para, image_map, image_idx_ref)
                    if md:
                        md_parts.append(md)
                    break

        elif tag == 'tbl':
            for table in doc.tables:
                if table._element is element:
                    md_parts.append(table_to_md(table))
                    break

    # Write output
    with open(output_md, 'w', encoding='utf-8') as f:
        f.write(''.join(md_parts))

    md_size = os.path.getsize(output_md)
    print(f'\nDone!')
    print(f'  Markdown: {output_md}')
    print(f'  Images:   {len(image_map)} extracted to {images_dir}')
    print(f'  Tables:   {len(doc.tables)}')
    print(f'  Size:     {md_size / 1024:.1f} KB')

    return output_md


def main():
    parser = argparse.ArgumentParser(
        description='Convert Word (.docx) to Markdown for Sloth-BookSmith-Den'
    )
    parser.add_argument('input', help='Input .docx file')
    parser.add_argument('-o', '--output', help='Output .md file (default: auto)')
    parser.add_argument('-d', '--output-dir', help='Output directory (default: same as input)')
    parser.add_argument('--images-dir', help='Images output directory (default: OUTPUT_DIR/images)')
    parser.add_argument('--title', help='Override book title')
    parser.add_argument('--author', help='Override author')
    parser.add_argument('--date', help='Override date (YYYY-MM-DD)')
    parser.add_argument('--subtitle', help='Override subtitle')
    args = parser.parse_args()

    convert(
        input_docx=args.input,
        output_md=args.output,
        output_dir=args.output_dir,
        images_dir=args.images_dir,
        title=args.title,
        author=args.author,
        date=args.date,
        subtitle=args.subtitle,
    )


if __name__ == '__main__':
    main()
